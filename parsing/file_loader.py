from pypdf import PdfReader
from docx import Document
import re
import os
from datetime import datetime

def load_pdf(file_path):
    reader = PdfReader(file_path)
    number_of_pages = len(reader.pages)
    text = ""
    for page in reader.pages:
        page_text =  page.extract_text()
        if page_text:
            text +=  page_text + "\n"
    return text

def load_docx(file_path):
    doc = Document(file_path)
    text = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)
    
    xml = doc.element.xml
    matches = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml)
    
    for match in matches:
        if match.strip():
            text.append(match)

    return "\n".join(dict.fromkeys(text))


def extract_text(file_path: str) -> str:
    extension = os.path.splitext(file_path)[1].lower()
    
    if extension == ".pdf":
        return load_pdf(file_path)
    elif extension == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Format not supported: {extension}")


def save_processed_text(
    text: str,
    source: str,
    original_file_name: str,
    base_dir: str = "../data/processed"
) -> str:
    """
    Save extracted text to disk for debugging and traceability.

    Args:
        text (str): Extracted raw text
        source (str): 'resume' or 'job_description'
        original_file_name (str): Original uploaded file name
        base_dir (str): Base directory to store processed files

    Returns:
        str: Path to saved file
    """

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = os.path.splitext(original_file_name)[0]
    filename = f"{safe_name}_{timestamp}.txt"

    target_dir = os.path.join(base_dir, f"{source}s")
    os.makedirs(target_dir, exist_ok=True)

    file_path = os.path.join(target_dir, filename)

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text)

    return file_path

# text = extract_text('../data/raw/CV-Jessy Gigato.pdf')
# print(text)
# save_processed_text(text,"resume","CV-Jessy Gigato.pdf")